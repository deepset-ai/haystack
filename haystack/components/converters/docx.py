# SPDX-FileCopyrightText: 2022-present deepset GmbH <info@deepset.ai>
#
# SPDX-License-Identifier: Apache-2.0

import csv
import io
import os
import warnings
from dataclasses import dataclass
from enum import Enum
from io import StringIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from haystack import Document, component, default_from_dict, default_to_dict, logging
from haystack.components.converters.utils import get_bytestream_from_source, normalize_metadata
from haystack.dataclasses import ByteStream
from haystack.lazy_imports import LazyImport

logger = logging.getLogger(__name__)

with LazyImport("Run 'pip install python-docx'") as docx_import:
    import docx
    from docx.document import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph


@dataclass
class DOCXMetadata:
    """
    Describes the metadata of Docx file.

    :param author: The author
    :param category: The category
    :param comments: The comments
    :param content_status: The content status
    :param created: The creation date (ISO formatted string)
    :param identifier: The identifier
    :param keywords: Available keywords
    :param language: The language of the document
    :param last_modified_by: User who last modified the document
    :param last_printed: The last printed date (ISO formatted string)
    :param modified: The last modification date (ISO formatted string)
    :param revision: The revision number
    :param subject: The subject
    :param title: The title
    :param version: The version
    """

    author: str
    category: str
    comments: str
    content_status: str
    created: Optional[str]
    identifier: str
    keywords: str
    language: str
    last_modified_by: str
    last_printed: Optional[str]
    modified: Optional[str]
    revision: int
    subject: str
    title: str
    version: str


class DOCXTableFormat(Enum):
    """
    Supported formats for storing DOCX tabular data in a Document.
    """

    MARKDOWN = "markdown"
    CSV = "csv"

    def __str__(self):
        return self.value

    @staticmethod
    def from_str(string: str) -> "DOCXTableFormat":
        """
        Convert a string to a DOCXTableFormat enum.
        """
        enum_map = {e.value: e for e in DOCXTableFormat}
        table_format = enum_map.get(string.lower())
        if table_format is None:
            msg = f"Unknown table format '{string}'. Supported formats are: {list(enum_map.keys())}"
            raise ValueError(msg)
        return table_format


@component
class DOCXToDocument:
    """
    Converts DOCX files to Documents.

    Uses `python-docx` library to convert the DOCX file to a document.
    This component does not preserve page breaks in the original document.

    Usage example:
    ```python
    from haystack.components.converters.docx import DOCXToDocument, DOCXTableFormat

    converter = DOCXToDocument(table_format=DOCXTableFormat.CSV)
    results = converter.run(sources=["sample.docx"], meta={"date_added": datetime.now().isoformat()})
    documents = results["documents"]
    print(documents[0].content)
    # 'This is a text from the DOCX file.'
    ```
    """

    def __init__(self, table_format: Union[str, DOCXTableFormat] = DOCXTableFormat.CSV, store_full_path: bool = True):
        """
        Create a DOCXToDocument component.

        :param table_format: The format for table output. Can be either DOCXTableFormat.MARKDOWN,
            DOCXTableFormat.CSV, "markdown", or "csv". Defaults to DOCXTableFormat.CSV.
        :param store_full_path:
            If True, the full path of the file is stored in the metadata of the document.
            If False, only the file name is stored.
        """
        docx_import.check()
        self.table_format = DOCXTableFormat.from_str(table_format) if isinstance(table_format, str) else table_format
        self.store_full_path = store_full_path

    def to_dict(self) -> Dict[str, Any]:
        """
        Serializes the component to a dictionary.

        :returns:
            Dictionary with serialized data.
        """
        return default_to_dict(self, table_format=str(self.table_format), store_full_path=self.store_full_path)

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "DOCXToDocument":
        """
        Deserializes the component from a dictionary.

        :param data:
            The dictionary to deserialize from.
        :returns:
            The deserialized component.
        """
        if "table_format" in data["init_parameters"]:
            data["init_parameters"]["table_format"] = DOCXTableFormat.from_str(data["init_parameters"]["table_format"])
        return default_from_dict(cls, data)

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path, ByteStream]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        """
        Converts DOCX files to Documents.

        :param sources:
            List of file paths or ByteStream objects.
        :param meta:
            Optional metadata to attach to the Documents.
            This value can be either a list of dictionaries or a single dictionary.
            If it's a single dictionary, its content is added to the metadata of all produced Documents.
            If it's a list, the length of the list must match the number of sources, because the two lists will
            be zipped.
            If `sources` contains ByteStream objects, their `meta` will be added to the output Documents.
        :returns:
            A dictionary with the following keys:
            - `documents`: Created Documents
        """
        documents = []
        meta_list = normalize_metadata(meta=meta, sources_count=len(sources))

        for source, metadata in zip(sources, meta_list):
            try:
                bytestream = get_bytestream_from_source(source)
            except Exception as e:
                logger.warning("Could not read {source}. Skipping it. Error: {error}", source=source, error=e)
                continue
            try:
                docx_document = docx.Document(io.BytesIO(bytestream.data))
                elements = self._extract_elements(docx_document)
                text = "\n".join(elements)
            except Exception as e:
                logger.warning(
                    "Could not read {source} and convert it to a DOCX Document, skipping. Error: {error}",
                    source=source,
                    error=e,
                )
                continue

            warnings.warn(
                "The `store_full_path` parameter defaults to True, storing full file paths in metadata. "
                "In the 2.9.0 release, the default value for `store_full_path` will change to False, "
                "storing only file names to improve privacy.",
                DeprecationWarning,
            )

            docx_metadata = self._get_docx_metadata(document=docx_document)
            merged_metadata = {**bytestream.meta, **metadata, "docx": docx_metadata}

            if not self.store_full_path and "file_path" in bytestream.meta:
                file_path = bytestream.meta.get("file_path")
                if file_path:  # Ensure the value is not None for pylint
                    merged_metadata["file_path"] = os.path.basename(file_path)

            document = Document(content=text, meta=merged_metadata)
            documents.append(document)

        return {"documents": documents}

    def _extract_elements(self, document: "DocxDocument") -> List[str]:
        """
        Extracts elements from a DOCX file.

        :param document: The DOCX Document object.
        :returns: List of strings (paragraph texts and table representations) with page breaks added as '\f' characters.
        """
        elements = []
        for element in document.element.body:
            if element.tag.endswith("p"):
                paragraph = Paragraph(element, document)
                if paragraph.contains_page_break:
                    para_text = self._process_paragraph_with_page_breaks(paragraph)
                else:
                    para_text = paragraph.text
                elements.append(para_text)
            elif element.tag.endswith("tbl"):
                table = docx.table.Table(element, document)
                table_str = (
                    self._table_to_markdown(table)
                    if self.table_format == DOCXTableFormat.MARKDOWN
                    else self._table_to_csv(table)
                )
                elements.append(table_str)

        return elements

    def _process_paragraph_with_page_breaks(self, paragraph: "Paragraph") -> str:
        """
        Processes a paragraph with page breaks.

        :param paragraph: The DOCX paragraph to process.
        :returns: A string with page breaks added as '\f' characters.
        """
        para_text = ""
        # Usually, just 1 page break exists, but could be more if paragraph is really long, so we loop over them
        for pb_index, page_break in enumerate(paragraph.rendered_page_breaks):
            # Can only extract text from first paragraph page break, unfortunately
            if pb_index == 0:
                if page_break.preceding_paragraph_fragment:
                    para_text += page_break.preceding_paragraph_fragment.text
                para_text += "\f"
                if page_break.following_paragraph_fragment:
                    # following_paragraph_fragment contains all text for remainder of paragraph.
                    # However, if the remainder of the paragraph spans multiple page breaks, it won't include
                    # those later page breaks so we have to add them at end of text in the `else` block below.
                    # This is not ideal, but this case should be very rare and this is likely good enough.
                    para_text += page_break.following_paragraph_fragment.text
            else:
                para_text += "\f"
        return para_text

    def _table_to_markdown(self, table: "Table") -> str:
        """
        Converts a DOCX table to a Markdown string.

        :param table: The DOCX table to convert.
        :returns: A Markdown string representation of the table.
        """
        markdown: List[str] = []
        max_col_widths: List[int] = []

        # Calculate max width for each column
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if i >= len(max_col_widths):
                    max_col_widths.append(len(cell_text))
                else:
                    max_col_widths[i] = max(max_col_widths[i], len(cell_text))

        # Process rows
        for i, row in enumerate(table.rows):
            md_row = [cell.text.strip().ljust(max_col_widths[j]) for j, cell in enumerate(row.cells)]
            markdown.append("| " + " | ".join(md_row) + " |")

            # Add separator after header row
            if i == 0:
                separator = ["-" * max_col_widths[j] for j in range(len(row.cells))]
                markdown.append("| " + " | ".join(separator) + " |")

        return "\n".join(markdown)

    def _table_to_csv(self, table: "Table") -> str:
        """
        Converts a DOCX table to a CSV string.

        :param table: The DOCX table to convert.
        :returns: A CSV string representation of the table.
        """
        csv_output = StringIO()
        csv_writer = csv.writer(csv_output, quoting=csv.QUOTE_MINIMAL)

        # Process rows
        for row in table.rows:
            csv_row = [cell.text.strip() for cell in row.cells]
            csv_writer.writerow(csv_row)

        # Get the CSV as a string and strip any trailing newlines
        csv_string = csv_output.getvalue().strip()
        csv_output.close()

        return csv_string

    def _get_docx_metadata(self, document: "DocxDocument") -> DOCXMetadata:
        """
        Get all relevant data from the 'core_properties' attribute from a DOCX Document.

        :param document:
            The DOCX Document you want to extract metadata from

        :returns:
            A `DOCXMetadata` dataclass all the relevant fields from the 'core_properties'
        """
        return DOCXMetadata(
            author=document.core_properties.author,
            category=document.core_properties.category,
            comments=document.core_properties.comments,
            content_status=document.core_properties.content_status,
            created=(document.core_properties.created.isoformat() if document.core_properties.created else None),
            identifier=document.core_properties.identifier,
            keywords=document.core_properties.keywords,
            language=document.core_properties.language,
            last_modified_by=document.core_properties.last_modified_by,
            last_printed=(
                document.core_properties.last_printed.isoformat() if document.core_properties.last_printed else None
            ),
            modified=(document.core_properties.modified.isoformat() if document.core_properties.modified else None),
            revision=document.core_properties.revision,
            subject=document.core_properties.subject,
            title=document.core_properties.title,
            version=document.core_properties.version,
        )
